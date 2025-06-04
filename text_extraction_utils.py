import os
import re
import logging
import hashlib
import mimetypes
from werkzeug.utils import secure_filename
from pathlib import Path
from datetime import datetime


try:
    import fitz  
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import docx  
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)

MAX_FILE_SIZE_MB = 50
MIN_CONTENT_LENGTH = 50
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200

ALLOWED_EXTENSIONS = {
    'pdf': {'pdf'},
    'text': {'txt', 'md', 'rtf'},
    'document': {'docx', 'doc'}
}

FILE_SIZE_LIMITS = {
    'pdf': 50 * 1024 * 1024,     
    'text': 10 * 1024 * 1024,    
    'document': 25 * 1024 * 1024  
}
def chunk_text(text, chunk_size=CHUNK_SIZE, overlap=CHUNK_OVERLAP):
    """Smart text chunking with overlap for better context"""
    if not text or not isinstance(text, str):
        return []
    
    text = text.strip()
    text_length = len(text)
    
    if text_length <= chunk_size:
        return [text] if text_length >= MIN_CONTENT_LENGTH else []
    
    chunks = []
    start = 0
    
    while start < text_length:
        end = start + chunk_size
        
        if end < text_length:
            # Smart boundary detection
            boundaries = [
                ('\n\n', 300),  # Paragraph boundaries
                ('.', 200),     # Sentence boundaries
                ('!', 200),     # Exclamation boundaries
                ('?', 200),     # Question boundaries
                (';', 150),     # Semicolon boundaries
                ('\n', 100),    # Line boundaries
                (' ', 50)       # Word boundaries
            ]
            
            best_boundary = end
            for boundary_char, search_range in boundaries:
                search_start = max(start, end - search_range)
                boundary_pos = text.rfind(boundary_char, search_start, end)
                if boundary_pos > start:
                    best_boundary = boundary_pos + len(boundary_char)
                    break
            
            end = best_boundary
        
        chunk = text[start:end].strip()
        
        # Validate chunk
        if len(chunk) >= MIN_CONTENT_LENGTH and len(chunk.split()) >= 3:
            chunks.append(chunk)
        
        start = max(end - overlap, start + 1)
        if start >= text_length:
            break
    
    return chunks

def detect_file_type(file):
    """Detect file type from filename and MIME type"""
    if not hasattr(file, 'filename') or not file.filename:
        return None
        
    filename = file.filename.lower()
    ext = filename.rsplit('.', 1)[1] if '.' in filename else ''
    
    # Check extension first
    for type_key, extensions in ALLOWED_EXTENSIONS.items():
        if ext in extensions:
            return type_key
    
    return None

def validate_file(file):
    """Validate file for processing"""
    if not file or not hasattr(file, 'filename'):
        return False, "Invalid file provided"
    
    if not file.filename:
        return False, "No filename provided"
    
    filename = file.filename
    if '..' in filename or '/' in filename or '\\' in filename:
        return False, "Invalid filename - security violation"
    
    file_type = detect_file_type(file)
    if not file_type:
        return False, f"Unsupported file type: {filename}"
    
    return True, file_type

def validate_file_size(file_path, file_type):
    """Validate file size based on type"""
    if not os.path.exists(file_path):
        return False, "File does not exist"
    
    file_size = os.path.getsize(file_path)
    limit = FILE_SIZE_LIMITS.get(file_type, MAX_FILE_SIZE_MB * 1024 * 1024)
    
    if file_size > limit:
        size_mb = file_size / (1024 * 1024)
        limit_mb = limit / (1024 * 1024)
        return False, f"File size ({size_mb:.1f}MB) exceeds limit ({limit_mb:.1f}MB)"
    
    return True, None

def extract_text_from_pdf(pdf_path):
    """Extract text from PDF using PyMuPDF"""
    if not PDF_AVAILABLE:
        return "", "PyMuPDF not available. Install with: pip install PyMuPDF"
    
    if not os.path.exists(pdf_path):
        return "", f"PDF file not found: {pdf_path}"
    
    try:
        # Validate file size
        size_ok, error = validate_file_size(pdf_path, 'pdf')
        if not size_ok:
            return "", error
        
        doc = fitz.open(pdf_path)
        text_parts = []
        total_pages = len(doc)
        
        if total_pages == 0:
            doc.close()
            return "", "PDF has no pages"
        
        successful_pages = 0
        
        for page_num in range(total_pages):
            try:
                page = doc[page_num]
                page_text = page.get_text()
                
                if page_text and page_text.strip():
                    cleaned_text = page_text.strip()
                    if len(cleaned_text) > 10:  # Minimum content threshold
                        text_parts.append(f"[Page {page_num + 1}]\n{cleaned_text}")
                        successful_pages += 1
                
            except Exception as e:
                logger.warning(f"Error processing page {page_num + 1}: {e}")
                continue
        
        doc.close()
        
        if successful_pages == 0:
            return "", "No readable content found in PDF"
        
        extracted_text = "\n\n".join(text_parts)
        logger.info(f"PDF extraction: {successful_pages}/{total_pages} pages, {len(extracted_text)} characters")
        
        return extracted_text, None
        
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        return "", f"PDF extraction failed: {str(e)}"

def extract_text_from_docx(docx_path):
    """Extract text from DOCX using python-docx"""
    if not DOCX_AVAILABLE:
        return "", "python-docx not available. Install with: pip install python-docx"
    
    if not os.path.exists(docx_path):
        return "", f"DOCX file not found: {docx_path}"
    
    try:
        # Validate file size
        size_ok, error = validate_file_size(docx_path, 'document')
        if not size_ok:
            return "", error
        
        doc = docx.Document(docx_path)
        text_parts = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        if not text_parts:
            return "", "No readable content found in DOCX"
        
        extracted_text = "\n\n".join(text_parts)
        logger.info(f"DOCX extraction: {len(text_parts)} elements, {len(extracted_text)} characters")
        
        return extracted_text, None
        
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        return "", f"DOCX extraction failed: {str(e)}"

def extract_text_from_text_file(file_path):
    """Extract text from plain text files with encoding detection"""
    if not os.path.exists(file_path):
        return "", f"Text file not found: {file_path}"
    
    try:
        # Validate file size
        size_ok, error = validate_file_size(file_path, 'text')
        if not size_ok:
            return "", error
        
        # Try multiple encodings
        encodings = ['utf-8', 'utf-16', 'utf-8-sig', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                
                if content and content.strip():
                    logger.info(f"Text file read with {encoding}: {len(content)} characters")
                    return content.strip(), None
                    
            except (UnicodeDecodeError, UnicodeError):
                continue
            except Exception as e:
                logger.error(f"Error reading with {encoding}: {e}")
                break
        
        return "", "Could not read text file with any supported encoding"
        
    except Exception as e:
        logger.error(f"Error processing text file: {e}")
        return "", f"Text file processing failed: {str(e)}"


import re
from typing import List, Dict, Any

def detect_headings_and_structure(text: str) -> List[Dict[str, Any]]:
    """
    Detect headings and create slide structure from text content
    Returns a list of slide dictionaries with content and metadata
    """
    lines = text.split('\n')
    headings = []
    
    # First pass: detect all potential headings
    for i, line in enumerate(lines):
        line = line.strip()
        if not line or len(line) < 3:
            continue
            
        heading_info = None
        
        # 1. Markdown headings (# ## ###)
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            title = line.lstrip('#').strip()
            if title and level <= 4:  # Only up to #### 
                heading_info = {
                    'line_number': i,
                    'level': level,
                    'title': title,
                    'type': 'markdown',
                    'confidence': 0.95
                }
        
        # 2. ALL CAPS lines (potential headings)
        elif (line.isupper() and 
              len(line.split()) <= 8 and 
              len(line.split()) >= 2 and
              not line.endswith('.') and
              not line.startswith('HTTP')):
            heading_info = {
                'line_number': i,
                'level': 2,
                'title': line.title(),
                'type': 'caps',
                'confidence': 0.7
            }
        
        # 3. Lines ending with colon (section headers)
        elif (line.endswith(':') and 
              len(line.split()) <= 10 and 
              len(line.split()) >= 2 and
              not line.lower().startswith('http')):
            heading_info = {
                'line_number': i,
                'level': 3,
                'title': line[:-1].strip(),
                'type': 'colon',
                'confidence': 0.6
            }
        
        # 4. Numbered sections (1. 2. etc.)
        elif re.match(r'^\d+\.\s+[A-Z]', line):
            heading_info = {
                'line_number': i,
                'level': 2,
                'title': re.sub(r'^\d+\.\s+', '', line),
                'type': 'numbered',
                'confidence': 0.8
            }
        
        # 5. Roman numerals (I. II. III.)
        elif re.match(r'^[IVX]+\.\s+[A-Z]', line):
            heading_info = {
                'line_number': i,
                'level': 2,
                'title': re.sub(r'^[IVX]+\.\s+', '', line),
                'type': 'roman',
                'confidence': 0.75
            }
        
        if heading_info:
            headings.append(heading_info)
    
    # Filter and validate headings
    valid_headings = filter_valid_headings(headings, lines)
    

    

def filter_valid_headings(headings: List[Dict], lines: List[str]) -> List[Dict]:
    """Filter out false positive headings"""
    if not headings:
        return []
    
    # Remove headings that are too close together (likely false positives)
    filtered = []
    for i, heading in enumerate(headings):
        if i == 0:
            filtered.append(heading)
            continue
        
        prev_heading = filtered[-1]
        line_diff = heading['line_number'] - prev_heading['line_number']
        
        # Skip if headings are too close (less than 3 lines apart)
        if line_diff < 3:
            # Keep the one with higher confidence
            if heading['confidence'] > prev_heading['confidence']:
                filtered[-1] = heading
            continue
        
        filtered.append(heading)
    
    return filtered




def detect_headings_and_structure_enhanced(text):
    """
    Enhanced heading detection with better accuracy and slide creation
    """
    lines = text.split('\n')
    headings = []
    
    # First pass: detect all potential headings with confidence scores
    for i, line in enumerate(lines):
        line = line.strip()
        if not line or len(line) < 3:
            continue
            
        heading_info = None
        
        # 1. Markdown headings (# ## ### ####)
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            title = line.lstrip('#').strip()
            if title and level <= 4:
                heading_info = {
                    'line_number': i,
                    'level': level,
                    'title': title,
                    'type': 'markdown',
                    'confidence': 0.95
                }
        
        # 2. ALL CAPS lines (section headers)
        elif (line.isupper() and 
              3 <= len(line.split()) <= 8 and 
              not line.endswith('.') and
              not line.startswith('HTTP') and
              not line.isdigit()):
            heading_info = {
                'line_number': i,
                'level': 2,
                'title': line.title(),
                'type': 'caps',
                'confidence': 0.75
            }
        
        # 3. Lines ending with colon (section headers)
        elif (line.endswith(':') and 
              3 <= len(line.split()) <= 10 and
              not line.lower().startswith('http') and
              not any(char.isdigit() for char in line[:3])):  # Not time stamps
            heading_info = {
                'line_number': i,
                'level': 3,
                'title': line[:-1].strip(),
                'type': 'colon',
                'confidence': 0.65
            }
        
        # 4. Numbered sections (1. 2. I. II. etc.)
        elif re.match(r'^\d+\.\s+[A-Z]', line):
            heading_info = {
                'line_number': i,
                'level': 2,
                'title': re.sub(r'^\d+\.\s+', '', line),
                'type': 'numbered',
                'confidence': 0.85
            }
        
        # 5. Roman numerals (I. II. III.)
        elif re.match(r'^[IVX]+\.\s+[A-Z]', line):
            heading_info = {
                'line_number': i,
                'level': 2,
                'title': re.sub(r'^[IVX]+\.\s+', '', line),
                'type': 'roman',
                'confidence': 0.8
            }
        
        # 6. Lines that look like titles (standalone, properly capitalized)
        elif (len(line.split()) >= 2 and 
              len(line.split()) <= 8 and
              line[0].isupper() and
              not line.endswith('.') and
              sum(1 for word in line.split() if word[0].isupper()) >= len(line.split()) * 0.7):
            # Check if next line is empty or content starts (not another potential heading)
            next_line = lines[i + 1].strip() if i + 1 < len(lines) else ""
            if not next_line or (next_line and not next_line[0].isupper()):
                heading_info = {
                    'line_number': i,
                    'level': 2,
                    'title': line,
                    'type': 'title_case',
                    'confidence': 0.6
                }
        
        if heading_info:
            headings.append(heading_info)
    
    # Filter out false positives
    valid_headings = filter_false_positives(headings, lines)
    
    if not valid_headings:
        # No clear headings found, use content-based sliding
        return create_slides_from_content_chunks(text)
    
    # Create slides from validated headings
    return create_slides_from_validated_headings(text, valid_headings)

def filter_false_positives(headings, lines):
    """Remove likely false positive headings"""
    if not headings:
        return []
    
    filtered = []
    
    # Remove headings that are too close together
    for i, heading in enumerate(headings):
        if i == 0:
            filtered.append(heading)
            continue
        
        prev_heading = filtered[-1]
        line_diff = heading['line_number'] - prev_heading['line_number']
        
        # Skip if headings are too close (less than 5 lines apart)
        if line_diff < 5:
            # Keep the one with higher confidence
            if heading['confidence'] > prev_heading['confidence']:
                filtered[-1] = heading
            continue
        
        # Skip if the content between headings is too short
        content_lines = lines[prev_heading['line_number']+1:heading['line_number']]
        content_text = '\n'.join(content_lines).strip()
        if len(content_text.split()) < 20:  # Less than 20 words
            if heading['confidence'] > prev_heading['confidence']:
                filtered[-1] = heading
            continue
        
        filtered.append(heading)
    
    # Remove headings with very low confidence if we have enough others
    if len(filtered) > 3:
        filtered = [h for h in filtered if h['confidence'] >= 0.6]
    
    return filtered

def create_slides_from_validated_headings(text, headings):
    """Create optimized slides from validated headings"""
    lines = text.split('\n')
    slides = []
    
    # Add title slide if document starts with substantial content before first heading
    if headings and headings[0]['line_number'] > 10:
        title_content = '\n'.join(lines[:headings[0]['line_number']]).strip()
        if len(title_content.split()) > 30:
            title = extract_title_from_content(title_content)
            slides.append({
                'title': title,
                'content': title_content,
                'layout': 'titleOnly',
                'is_title_slide': True,
                'confidence': 1.0,
                'source_lines': (0, headings[0]['line_number'])
            })
    
    # Create slides from headings
    for i, heading in enumerate(headings):
        start_line = heading['line_number']
        end_line = headings[i + 1]['line_number'] if i + 1 < len(headings) else len(lines)
        
        # Extract content between this heading and the next
        content_lines = lines[start_line + 1:end_line]
        content = '\n'.join(content_lines).strip()
        
        # Skip if content is too short
        if len(content.split()) < 15:
            continue
        
        # Determine the best layout for this content
        layout = determine_layout_from_content_enhanced(content, heading['level'])
        
        slides.append({
            'title': heading['title'],
            'content': content,
            'layout': layout,
            'heading_type': heading['type'],
            'level': heading['level'],
            'confidence': heading['confidence'],
            'is_heading': True,
            'source_lines': (start_line, end_line),
            'word_count': len(content.split())
        })
    
    # Add conclusion slide if there's substantial content after the last heading
    if headings and slides:
        last_heading_end = headings[-1]['line_number']
        last_slide_end = slides[-1]['source_lines'][1]
        remaining_content = '\n'.join(lines[last_slide_end:]).strip()
        
        if len(remaining_content.split()) > 30:
            slides.append({
                'title': 'Summary and Conclusions',
                'content': remaining_content,
                'layout': 'conclusion',
                'is_conclusion': True,
                'confidence': 0.8,
                'source_lines': (last_slide_end, len(lines))
            })
    
    # Ensure we have a reasonable number of slides (3-10)
    slides = optimize_slide_count(slides, text)
    
    return slides

def create_slides_from_content_chunks(text):
    """Fallback: create slides when no clear headings are found"""
    # Split by double line breaks first
    sections = [section.strip() for section in text.split('\n\n') if section.strip()]
    
    if len(sections) <= 2:
        # Very little structure, split by word count
        return create_slides_by_word_count(text)
    
    slides = []
    
    # First section as title slide if it's short
    if sections and len(sections[0].split()) < 50:
        title = extract_title_from_content(sections[0])
        slides.append({
            'title': title,
            'content': sections[0],
            'layout': 'titleOnly',
            'is_title_slide': True,
            'auto_generated': True
        })
        sections = sections[1:]
    
    # Convert remaining sections to slides
    for i, section in enumerate(sections):
        if len(section.split()) < 15:  # Too short, skip
            continue
            
        title = extract_title_from_section(section, i + 1)
        layout = determine_layout_from_content_enhanced(section, 2)
        
        slides.append({
            'title': title,
            'content': section,
            'layout': layout,
            'section_index': i + 1,
            'auto_generated': True
        })
    
    return slides

def create_slides_by_word_count(text):
    """Create slides by splitting based on word count when no structure is found"""
    words = text.split()
    total_words = len(words)
    
    # Aim for 150-300 words per slide
    target_words_per_slide = 200
    estimated_slides = max(3, min(8, total_words // target_words_per_slide))
    
    slides = []
    words_per_slide = total_words // estimated_slides
    
    for i in range(estimated_slides):
        start_word = i * words_per_slide
        end_word = (i + 1) * words_per_slide if i < estimated_slides - 1 else total_words
        
        slide_words = words[start_word:end_word]
        slide_content = ' '.join(slide_words)
        
        # Try to end at sentence boundary
        if i < estimated_slides - 1:  # Not the last slide
            last_sentences = slide_content.split('.')
            if len(last_sentences) > 1:
                slide_content = '.'.join(last_sentences[:-1]) + '.'
                # Add the remaining to next slide's start
                remaining = last_sentences[-1]
                if i + 1 < estimated_slides and remaining.strip():
                    words[start_word + words_per_slide:start_word + words_per_slide] = remaining.split()
        
        title = f"Key Points {i + 1}" if i > 0 else extract_title_from_content(slide_content)
        layout = 'titleOnly' if i == 0 else 'titleAndBullets'
        
        slides.append({
            'title': title,
            'content': slide_content,
            'layout': layout,
            'word_based': True,
            'slide_index': i + 1
        })
    
    return slides

def extract_title_from_section(content, section_number):
    """Extract a meaningful title from a content section"""
    lines = [line.strip() for line in content.split('\n') if line.strip()]
    
    if not lines:
        return f'Section {section_number}'
    
    # Use first line if it's short and looks like a title
    first_line = lines[0]
    if len(first_line.split()) <= 8 and not first_line.endswith('.'):
        return first_line
    
    # Extract key phrases from first sentence
    first_sentence = content.split('.')[0]
    words = first_sentence.split()[:6]
    
    # Clean up common starting words
    while words and words[0].lower() in ['the', 'this', 'that', 'these', 'those', 'a', 'an']:
        words = words[1:]
    
    if words:
        title = ' '.join(words)
        return title + ('...' if len(first_sentence.split()) > 6 else '')
    
    return f'Section {section_number}'

def optimize_slide_count(slides, original_text):
    """Optimize the number of slides to be between 3-10"""
    if not slides:
        return []
    
    # If too few slides, try to split large ones
    if len(slides) < 3:
        return expand_slides(slides, original_text)
    
    # If too many slides, merge smaller ones
    if len(slides) > 10:
        return consolidate_slides(slides)
    
    return slides

def expand_slides(slides, original_text):
    """Expand slides when we have too few"""
    # If we only have 1-2 slides, use word-based splitting
    if len(slides) <= 2:
        return create_slides_by_word_count(original_text)
    
    return slides

def consolidate_slides(slides):
    """Consolidate slides when we have too many"""
    # Sort by word count and merge smallest ones
    slides_with_size = [(slide, len(slide['content'].split())) for slide in slides]
    slides_with_size.sort(key=lambda x: x[1])  # Sort by word count
    
    consolidated = []
    i = 0
    
    while i < len(slides_with_size):
        current_slide, current_size = slides_with_size[i]
        
        # If this slide is small and we have too many slides, try to merge
        if current_size < 50 and len(slides_with_size) > 8 and i + 1 < len(slides_with_size):
            next_slide, next_size = slides_with_size[i + 1]
            
            # Merge if combined size is reasonable
            if current_size + next_size < 300:
                merged_slide = {
                    'title': current_slide['title'],
                    'content': current_slide['content'] + '\n\n' + next_slide['content'],
                    'layout': 'imageAndParagraph',  # Good for merged content
                    'merged': True,
                    'original_titles': [current_slide['title'], next_slide['title']]
                }
                consolidated.append(merged_slide)
                i += 2  # Skip next slide as it's been merged
                continue
        
        consolidated.append(current_slide)
        i += 1
    
    return consolidated[:10]  # Ensure we don't exceed 10 slides

def extract_title_from_content(content):
    """Extract a title from content block - enhanced version"""
    if not content:
        return 'Untitled'
    
    lines = [line.strip() for line in content.split('\n') if line.strip()]
    
    if not lines:
        return 'Untitled'
    
    # Use first line if it's short and looks like a title
    first_line = lines[0]
    if len(first_line.split()) <= 8 and not first_line.endswith('.'):
        # Clean up common prefixes
        cleaned = re.sub(r'^(chapter|section|part)\s+\d+:?\s*', '', first_line, flags=re.IGNORECASE)
        return cleaned if cleaned else first_line
    
    # Extract first few meaningful words
    first_sentence = content.split('.')[0]
    words = first_sentence.split()
    
    # Remove common starting words
    skip_words = ['the', 'this', 'that', 'these', 'those', 'a', 'an', 'in', 'on', 'at', 'for', 'with']
    while words and words[0].lower() in skip_words:
        words = words[1:]
    
    # Take first 4-6 meaningful words
    if words:
        title_words = words[:6]
        title = ' '.join(title_words)
        
        # Add ellipsis if we truncated
        if len(words) > 6:
            title += '...'
        
        # Capitalize properly
        return title.title()
    
    return 'Content Overview'

def extract_title_from_content(content: str) -> str:
    """Extract a title from content block"""
    lines = [line.strip() for line in content.split('\n') if line.strip()]
    
    if not lines:
        return 'Untitled'
    
    # Use first line if it's short and looks like a title
    first_line = lines[0]
    if len(first_line.split()) <= 8 and not first_line.endswith('.'):
        return first_line
    
    # Extract first few words
    words = content.split()[:6]
    return ' '.join(words) + ('...' if len(content.split()) > 6 else '')



def extract_text_from_file(file_path, file_type=None):
    """Universal text extraction based on file type"""
    if not file_type:
        # Auto-detect based on extension
        ext = file_path.lower().rsplit('.', 1)[-1] if '.' in file_path else ''
        
        if ext == 'pdf':
            file_type = 'pdf'
        elif ext in ['docx', 'doc']:
            file_type = 'document'
        else:
            file_type = 'text'
    
    # Route to appropriate extractor
    if file_type == 'pdf':
        return extract_text_from_pdf(file_path)
    elif file_type == 'document':
        return extract_text_from_docx(file_path)
    else:
        return extract_text_from_text_file(file_path)


def analyze_text_content(text):
    """Analyze text content and suggest slide count"""
    if not text:
        return {
            'chars': 0,
            'words': 0,
            'sentences': 0,
            'paragraphs': 0,
            'suggested_slides': 3
        }
    
    # Basic metrics
    chars = len(text)
    words = len(text.split()) if text.strip() else 0
    sentences = len(re.findall(r'[.!?]+', text))
    paragraphs = len([p for p in text.split('\n\n') if p.strip()])
    
    # Smart slide suggestion
    if words < 100:
        suggested_slides = 3
    elif words < 300:
        suggested_slides = 4
    elif words < 600:
        suggested_slides = 5
    elif words < 1000:
        suggested_slides = 6
    elif words < 1500:
        suggested_slides = 7
    elif words < 2500:
        suggested_slides = 8
    else:
        suggested_slides = min(10, max(8, words // 300))
    
    return {
        'chars': chars,
        'words': words,
        'sentences': sentences,
        'paragraphs': paragraphs,
        'suggested_slides': suggested_slides,
        'readability': 'high' if words > 500 else 'medium' if words > 100 else 'low'
    }

def process_uploaded_file(file, temp_dir):
    """Process uploaded file and extract text content"""
    try:
        # Validate file
        is_valid, result = validate_file(file)
        if not is_valid:
            return None, result
        
        file_type = result
        
        # Save file securely
        filename = secure_filename(file.filename)
        if not filename:
            return None, "Invalid filename"
        
        # Create unique filename to avoid conflicts
        file_hash = hashlib.md5(filename.encode()).hexdigest()[:8]
        safe_filename = f"{file_hash}_{filename}"
        file_path = os.path.join(temp_dir, safe_filename)
        
        # Ensure temp directory exists
        os.makedirs(temp_dir, exist_ok=True)
        
        # Save uploaded file
        file.save(file_path)
        
        # Extract text
        extracted_text, error = extract_text_from_file(file_path, file_type)
        
        # Clean up
        try:
            os.remove(file_path)
        except:
            pass
        
        if error:
            return None, error
        
        if not extracted_text or len(extracted_text.strip()) < MIN_CONTENT_LENGTH:
            return None, "Could not extract meaningful text from file"
        
        # Analyze content
        analysis = analyze_text_content(extracted_text)
                # Log extracted text to a file
        try:
            log_dir = os.path.join(temp_dir, "logs")
            os.makedirs(log_dir, exist_ok=True)
            
            log_filename = os.path.join(log_dir, "extracted_text.log")
            with open(log_filename, "a", encoding="utf-8") as log_file:
                log_file.write(f"\n----- {datetime.now()} - {filename} -----\n")
                log_file.write(extracted_text[:5000])  # Limit to first 5000 chars to avoid overload
                log_file.write("\n\n")
        except Exception as log_error:
            logger.warning(f"Failed to log extracted text: {log_error}")

        return {
            'text': extracted_text,
            'analysis': analysis,
            'file_type': file_type,
            'original_filename': filename
        }, None
        
    except Exception as e:
        logger.error(f"Error processing uploaded file: {e}")
        return None, f"File processing failed: {str(e)}"
